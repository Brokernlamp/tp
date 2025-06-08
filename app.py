import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import streamlit as st
import base64
from io import BytesIO
from docx import Document
from docx.shared import Inches
import requests
import json
import time
import streamlit as st

# Custom CSS for a modern and engaging UI
st.markdown(
    """
    <style>
    /* Main background and font */
    .stApp {
        background: linear-gradient(135deg, #1e3c72, #2a5298);
        color: white;
        font-family: 'Arial', sans-serif;
    }

    /* Title styling */
    h1 {
        text-align: center;
        font-weight: bold;
        color: #ff9f43;
        font-size: 3rem;
        text-shadow: 2px 2px 4px rgba(0, 0, 0, 0.5);
        margin-bottom: 20px;
    }

    /* Sidebar styling */
    .css-1d391kg {
        background-color: rgba(255, 255, 255, 0.1);
        border-radius: 10px;
        padding: 20px;
        margin: 10px;
    }

    /* Button styling */
    .stButton button {
        background-color: #ff9f43;
        color: white;
        border-radius: 8px;
        padding: 10px 20px;
        font-size: 1rem;
        font-weight: bold;
        border: none;
        transition: background-color 0.3s ease;
    }

    .stButton button:hover {
        background-color: #ff7f00;
    }

    /* File uploader styling */
    .stFileUploader {
        background-color: rgba(255, 255, 255, 0.1);
        border-radius: 10px;
        padding: 20px;
        margin: 10px;
    }

    /* Logo styling */
    .logo {
        position: fixed;
        top: 60px;
        left: 20px;
        width: 80px;
        z-index: 100;
        border-radius: 50%;
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.2);
    }

    /* Gradient border for plots */
    .stPlot {
        border: 4px solid;
        border-image: linear-gradient(135deg, #ff9f43, #ff7f00) 1;
        border-radius: 10px;
        padding: 10px;
    }

    /* AI recommendation styling */
    .ai-recommendation {
        background-color: rgba(255, 255, 255, 0.15);
        border-left: 4px solid #ff9f43;
        padding: 15px;
        margin: 15px 0;
        border-radius: 5px;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# Function to add background image
def add_bg_from_local(image_file):
    try:
        with open(image_file, "rb") as f:
            base64_image = base64.b64encode(f.read()).decode()

        bg_style = f"""
        <style>
        .stApp {{
            background-image: url("data:image/png;base64,{base64_image}");
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
        }}
        </style>
        """
        st.markdown(bg_style, unsafe_allow_html=True)
    except FileNotFoundError:
        pass

# Function to add a fixed logo at the top-left
def add_logo(logo_path, width=80):
    try:
        with open(logo_path, "rb") as f:
            base64_logo = base64.b64encode(f.read()).decode()

        logo_style = f"""
        <style>
        .logo {{
            position: fixed;
            top: 60px;
            left: 20px;
            width: {width}px;
        }}
        </style>
        <img src="data:image/png;base64,{base64_logo}" class="logo">
        """
        st.markdown(logo_style, unsafe_allow_html=True)
    except FileNotFoundError:
        pass

# Try to add background and logo
try:
    add_bg_from_local("image.png")
    add_logo("logo.png", width=80)
except Exception:
    pass

# Title with emoji and shadow
st.markdown(
    "<h1>🏆 Player Performance Dashboard</h1>",
    unsafe_allow_html=True,
)

# Folder paths
reaction_time_folder_path = r"Reaction Time"
agility_folder_path = r"Agility"
decision_making_folder_path = r"Decision Making"
stamina_folder_path = r"Stamina"

# Hugging Face API integration
def get_ai_recommendations(player_data, metric_type, selected_metric, percentile, comparison_quality):
    import requests
    import json

    api_key = st.secrets["OPENROUTER_API_KEY"] # Replace with your real OpenRouter key

    prompt = f"""
A player is in the {percentile:.1f}th percentile for {metric_type} ({selected_metric}), which is considered {comparison_quality}.
As a top-level sports coach, suggest:
1. Name the excersise that international athelete would use to improve in {metric_type} 
2. One line explanation on why to use the suggested drill
3. Short brief on number of sets and correct way of the drill
Remember to keep the text format simple and prefer new line for better representation.


Avoid generic advice. Be creative, clear, and motivational.
"""

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }

    url = "https://openrouter.ai/api/v1/chat/completions"

    payload = {
        "model": "anthropic/claude-sonnet-4",
        "messages": [
            {"role": "system", "content": "You are a professional sports coach giving personalized athlete improvement tips."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.8,
        "max_tokens": 500
    }

    try:
        response = requests.post(url, headers=headers, json=payload)

        if response.status_code == 200:
            result = response.json()
            return result["choices"][0]["message"]["content"]
        else:
            return f"⚠️ API Error ({response.status_code}): {response.text}"

    except Exception as e:
        return f"❌ Failed to get AI recommendation: {e}"

    try:
        # Initialize with a default recommendation in case the API call fails
        recommendation = "Based on your performance data, focus on consistency and technique. Consider working with a trainer to develop a personalized improvement plan."
        
        # Make the API request
        response = requests.post(api_url, headers=headers, data=json.dumps(payload))
        
        if response.status_code == 200:
            result = response.json()
            if isinstance(result, list) and len(result) > 0 and "generated_text" in result[0]:
                recommendation = result[0]["generated_text"]
            elif isinstance(result, dict) and "generated_text" in result:
                recommendation = result["generated_text"]
        
        # Add fallback for API rate limiting
        if "error" in response.text and "rate limit" in response.text.lower():
            st.warning("API rate limit reached. Using default recommendations.")
            time.sleep(2)  # Wait before retrying
            
        return recommendation
    
    except Exception as e:
        st.error(f"Error getting AI recommendations: {e}")
        # Return a generic recommendation if the API call fails
        return f"""
        Based on your {metric_type} performance in the {selected_metric} metric (at the {percentile:.1f}th percentile), 
        I recommend focusing on consistent training routines that target this specific area. 
        Consider exercises that progressively challenge your current ability level. 
        Work with your coach to develop more specialized drills.

        """

# Enhanced function to calculate multiple metrics for each file in the folder
def calculate_metrics(folder_path, metric_type, column_name='TapTS', output_folder=None):
    results = []
    player_count = 1

    for file_name in os.listdir(folder_path):
        if file_name.endswith(".xlsx") and not file_name.startswith("~$"):
            file_path = os.path.join(folder_path, file_name)
            data = pd.read_excel(file_path)

            if column_name in data.columns:
                player_metrics = {"Player": f'Player {player_count}'}
                
                if metric_type == "Reaction Time":
                    # Primary metric - Simple Average
                    player_metrics["Average RT"] = data[column_name].mean()
                    
                    # Alternative metrics for Reaction Time
                    player_metrics["Median RT"] = data[column_name].median()
                    player_metrics["Min RT"] = data[column_name].min()
                    player_metrics["Consistency (Std Dev)"] = data[column_name].std()
                    
                    # Calculate improvement (difference between first half and second half)
                    midpoint = len(data) // 2
                    first_half = data[column_name].iloc[:midpoint].mean()
                    second_half = data[column_name].iloc[midpoint:].mean()
                    player_metrics["Improvement"] = first_half - second_half
                    
                elif metric_type == "Agility":
                    # Primary metric - Average difference between consecutive points
                    player_metrics["Average Change"] = data[column_name].diff().abs().mean()
                    
                    # Alternative metrics for Agility
                    player_metrics["Max Change"] = data[column_name].diff().abs().max()
                    player_metrics["Change Consistency"] = data[column_name].diff().abs().std()
                    
                    # Calculate directional changes (how often direction changes)
                    diffs = data[column_name].diff()
                    direction_changes = ((diffs[1:] * diffs[:-1].values) < 0).sum()
                    player_metrics["Direction Changes"] = direction_changes
                    
                    # Calculate agility ratio (direction changes / total movements)
                    player_metrics["Agility Ratio"] = direction_changes / (len(data) - 1) if len(data) > 1 else 0
                    
                elif metric_type == "Decision Making":
                    # Decision making metrics
                    if 'Correct' in data.columns:
                        player_metrics["Accuracy"] = data['Correct'].mean() * 100
                    player_metrics["Response Time"] = data[column_name].mean()
                    player_metrics["Decision Consistency"] = data[column_name].std()
                    
                elif metric_type == "Stamina":
                    # Stamina metrics
                    player_metrics["Total Time"] = data[column_name].sum()
                    
                    # Calculate endurance factor (ratio of last quarter to first quarter times)
                    quarter_point = len(data) // 4
                    first_quarter = data[column_name].iloc[:quarter_point].mean()
                    last_quarter = data[column_name].iloc[-quarter_point:].mean()
                    player_metrics["Endurance Factor"] = last_quarter / first_quarter if first_quarter > 0 else 0
                    
                results.append(player_metrics)
                player_count += 1

    # Convert to DataFrame
    results_df = pd.DataFrame(results)

    if output_folder and not results_df.empty:
        os.makedirs(output_folder, exist_ok=True)
        output_file_path = os.path.join(output_folder, f'{metric_type.lower().replace(" ", "_")}_metrics.xlsx')
        results_df.to_excel(output_file_path, index=False)
        return results_df, output_file_path

    return results_df, None

# Function to handle metrics with multiple calculation methods
def handle_metric(folder_path, metric_type):
    st.header(f"{metric_type} Analysis")
    
    # Add subheader explaining multiple metrics
    st.subheader("Multiple Metrics Available")
    
    if metric_type == "Reaction Time":
        st.markdown("""
        Select the metric that best suits your analysis:
        - **Average RT**: Standard average reaction time
        - **Median RT**: Middle value (less affected by outliers)
        - **Min RT**: Fastest reaction time (peak performance)
        - **Consistency**: How consistent the reactions are (lower is better)
        - **Improvement**: Difference between first and second half (positive means improvement)
        """)
        
        metric_option = st.selectbox(
            "Select Reaction Time Metric",
            ["Average RT", "Median RT", "Min RT", "Consistency (Std Dev)", "Improvement"]
        )
        
    elif metric_type == "Agility":
        st.markdown("""
        Select the metric that best suits your analysis:
        - **Average Change**: Standard measure of movement speed
        - **Max Change**: Fastest movement recorded
        - **Change Consistency**: How consistent the movements are
        - **Direction Changes**: How many times direction changed
        - **Agility Ratio**: Proportion of movements that changed direction
        """)
        
        metric_option = st.selectbox(
            "Select Agility Metric",
            ["Average Change", "Max Change", "Change Consistency", "Direction Changes", "Agility Ratio"]
        )
    else:
        metric_option = "Primary"

    if st.button(f"Create {metric_type} Metrics File"):
        metrics_df, output_file_path = calculate_metrics(folder_path, metric_type, output_folder=folder_path)
        if output_file_path:
            st.success(f"{metric_type} metrics file created at {output_file_path}")
            st.dataframe(metrics_df)
        else:
            st.error(f"Failed to create {metric_type} metrics file")

    uploaded_file = st.file_uploader(f"Upload New Player's {metric_type} Data", type="xlsx")
    if uploaded_file is not None:
        analyze_uploaded_file(uploaded_file, metric_type, folder_path, selected_metric=metric_option)

# Function to analyze uploaded file with multiple metric options
def analyze_uploaded_file(uploaded_file, metric_type, folder_path, selected_metric="Primary"):
    try:
        new_player_data = pd.read_excel(uploaded_file)

        if 'TapTS' in new_player_data.columns:
            # Calculate metrics based on the selected metric type
            if metric_type == "Reaction Time":
                if selected_metric == "Average RT":
                    new_player_metric = new_player_data['TapTS'].mean()
                    metric_description = "Average Reaction Time"
                elif selected_metric == "Median RT":
                    new_player_metric = new_player_data['TapTS'].median()
                    metric_description = "Median Reaction Time"
                elif selected_metric == "Min RT":
                    new_player_metric = new_player_data['TapTS'].min()
                    metric_description = "Minimum Reaction Time"
                elif selected_metric == "Consistency (Std Dev)":
                    new_player_metric = new_player_data['TapTS'].std()
                    metric_description = "Reaction Time Consistency (lower is better)"
                elif selected_metric == "Improvement":
                    midpoint = len(new_player_data) // 2
                    first_half = new_player_data['TapTS'].iloc[:midpoint].mean()
                    second_half = new_player_data['TapTS'].iloc[midpoint:].mean()
                    new_player_metric = first_half - second_half
                    metric_description = "Reaction Time Improvement (positive is better)"
                else:
                    new_player_metric = new_player_data['TapTS'].mean()
                    metric_description = "Average Reaction Time"
                    
            elif metric_type == "Agility":
                if selected_metric == "Average Change":
                    new_player_metric = new_player_data['TapTS'].diff().abs().mean()
                    metric_description = "Average Movement Speed"
                elif selected_metric == "Max Change":
                    new_player_metric = new_player_data['TapTS'].diff().abs().max()
                    metric_description = "Maximum Movement Speed"
                elif selected_metric == "Change Consistency":
                    new_player_metric = new_player_data['TapTS'].diff().abs().std()
                    metric_description = "Movement Consistency (lower is better)"
                elif selected_metric == "Direction Changes":
                    diffs = new_player_data['TapTS'].diff()
                    new_player_metric = ((diffs[1:] * diffs[:-1].values) < 0).sum()
                    metric_description = "Number of Direction Changes"
                elif selected_metric == "Agility Ratio":
                    diffs = new_player_data['TapTS'].diff()
                    direction_changes = ((diffs[1:] * diffs[:-1].values) < 0).sum()
                    new_player_metric = direction_changes / (len(new_player_data) - 1) if len(new_player_data) > 1 else 0
                    metric_description = "Agility Ratio (higher is better)"
                else:
                    new_player_metric = new_player_data['TapTS'].diff().abs().mean()
                    metric_description = "Average Movement Speed"
            else:
                # Generic fallback
                new_player_metric = new_player_data['TapTS'].mean()
                metric_description = f"{metric_type} Metric"

            # Load comparison data
            metric_file = f"{metric_type.lower().replace(' ', '_')}_metrics.xlsx"
            avg_file_path = os.path.join(folder_path, metric_file)
            
            if os.path.exists(avg_file_path):
                avg_df = pd.read_excel(avg_file_path)
                
                # Display the metrics file for reference
                with st.expander("Show All Player Metrics"):
                    st.dataframe(avg_df)
                
                # Extract the column for comparison
                if selected_metric in avg_df.columns:
                    comparison_column = selected_metric
                elif metric_type == "Reaction Time" and "Average RT" in avg_df.columns:
                    comparison_column = "Average RT"
                elif metric_type == "Agility" and "Average Change" in avg_df.columns:
                    comparison_column = "Average Change"
                elif "Metric Value" in avg_df.columns:
                    comparison_column = "Metric Value"
                else:
                    # If column not found, use the second column (assuming first is Player name)
                    comparison_column = avg_df.columns[1]
                
                overall_avg = avg_df[comparison_column].mean()
                
                # For some metrics, lower is better (consistency, reaction time)
                lower_is_better = any(term in selected_metric.lower() for term in ["consistency", "average rt", "median rt"])
                if lower_is_better:
                    percentile = (np.sum(avg_df[comparison_column] > new_player_metric) / len(avg_df)) * 100
                    comparison = "lower" if new_player_metric < overall_avg else "higher"
                    comparison_quality = "better" if new_player_metric < overall_avg else "needs improvement"
                else:
                    percentile = (np.sum(avg_df[comparison_column] < new_player_metric) / len(avg_df)) * 100
                    comparison = "higher" if new_player_metric > overall_avg else "lower"
                    comparison_quality = "better" if new_player_metric > overall_avg else "needs improvement"

                st.write(f"### {selected_metric} Analysis")
                summary_table = pd.DataFrame({
                    "Metric": [f"Group Average", f"Your {metric_description}", "Percentile"],
                    "Value": [f"{overall_avg:.2f}", f"{new_player_metric:.2f}", f"{percentile:.2f}%"]
                })
                st.table(summary_table)
                
                st.info(f"Your {metric_description} is {comparison} than the group average, which is {comparison_quality}.")
                
                # Create player comparison visualization
                fig, ax = plt.subplots(figsize=(12, 8))
                
                # Plot histogram of all players
                ax.hist(avg_df[comparison_column], bins=10, alpha=0.7, color='blue', label='All Players')
                
                # Add line for the new player
                ax.axvline(x=new_player_metric, color='red', linestyle='--', linewidth=2,
                           label=f'Your {metric_description}')
                
                # Add group average line
                ax.axvline(x=overall_avg, color='green', linestyle=':', linewidth=2,
                          label='Group Average')
                
                ax.set_xlabel(metric_description, fontsize=12)
                ax.set_ylabel("Number of Players", fontsize=12)
                ax.set_title(f"{metric_type}: {selected_metric} Distribution", fontsize=16)
                ax.legend()
                ax.grid(True, alpha=0.3)
                
                st.pyplot(fig)
                
                # Alternative plot: percentile position
                plot1 = plot_player_comparison(new_player_metric, 
                                             pd.DataFrame({'Player': avg_df['Player'], 
                                                           'Metric Value': avg_df[comparison_column]}),
                                             f"{metric_type}: {selected_metric}")

                # Get AI recommendations
                st.subheader("🤖 AI Coach Recommendations")
                
                with st.spinner("Generating personalized AI recommendations..."):
                    ai_recommendation = get_ai_recommendations(
                        new_player_data, 
                        metric_type, 
                        selected_metric, 
                        percentile, 
                        comparison_quality
                    )
                
                # Display AI recommendations with custom styling
                st.markdown(
                    f"""<div class="ai-recommendation">
                    <h4>Personalized Training Recommendations</h4>
                    {ai_recommendation}
                    </div>""", 
                    unsafe_allow_html=True
                )

                # Create Word document with report and AI recommendations
                doc_io = create_word_document_enhanced(summary_table, plot1, 
                                                     metric_type, selected_metric, 
                                                     new_player_metric, overall_avg,
                                                     percentile, comparison_quality,
                                                     ai_recommendation)

                # Provide download button
                st.download_button(
                    label="Download Report with AI Recommendations",
                    data=doc_io,
                    file_name=f"{metric_type.lower().replace(' ', '_')}_{selected_metric.lower().replace(' ', '_')}_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning(f"Could not find comparison data at {avg_file_path}. Please create the metrics file first.")
        else:
            st.error("The uploaded file is missing the 'TapTS' column.")
    except Exception as e:
        st.error(f"An error occurred: {e}")
        st.exception(e)

# Function to create an enhanced Word document with the outputs
def create_word_document_enhanced(summary_table, plot1, metric_type, selected_metric, 
                                 player_value, average_value, percentile, comparison,
                                 ai_recommendation=""):
    doc = Document()
    doc.add_heading('Player Performance Report', 0)
    
    # Add detailed header
    doc.add_heading(f'{metric_type} Analysis: {selected_metric}', 1)
    
    # Add summary paragraph
    doc.add_paragraph(f"This report analyzes the player's {metric_type.lower()} using the {selected_metric} metric. "
                     f"The player's value of {player_value:.2f} is compared against the group average of {average_value:.2f}. "
                     f"The player ranks at the {percentile:.2f}th percentile, which is {comparison}.")
    
    # Add summary table
    doc.add_heading('Summary Table', level=2)
    table = doc.add_table(rows=1, cols=len(summary_table.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(summary_table.columns):
        hdr_cells[i].text = col

    for _, row in summary_table.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    # Add interpretation section
    doc.add_heading('Interpretation', level=2)
    
    if metric_type == "Reaction Time":
        doc.add_paragraph(f"The {selected_metric} is a measure of how quickly the player responds to stimuli. "
                         f"{'Lower values indicate better performance in terms of speed and reflexes.' if 'Consistency' not in selected_metric else 'Lower values indicate more consistent performance.'}")
    elif metric_type == "Agility":
        doc.add_paragraph(f"The {selected_metric} measures the player's ability to change direction quickly and efficiently. "
                         f"{'Higher values generally indicate better agility and movement capacity.' if not any(x in selected_metric for x in ['Consistency']) else 'Lower values indicate more consistent movement patterns.'}")
    
    # Add AI-powered recommendations section
    doc.add_heading('AI Coach Recommendations', level=2)
    
    # Style the AI recommendation text
    ai_paragraph = doc.add_paragraph()
    ai_run = ai_paragraph.add_run(ai_recommendation if ai_recommendation else 
                               "Based on your performance data, we recommend focusing on consistency and technique. "
                               "Consider working with a coach on personalized drills to improve this specific aspect of your performance.")
    
    # Add recommendations section
    doc.add_heading('General Recommendations', level=2)
    if percentile > 75:
        doc.add_paragraph("The player shows excellent performance in this metric. Continue current training methods and consider more advanced drills to maintain this advantage.")
    elif percentile > 50:
        doc.add_paragraph("The player shows above-average performance. Focus on consistency and specialized drills to reach the top tier.")
    elif percentile > 25:
        doc.add_paragraph("The player shows moderate performance. Targeted training focusing specifically on this aspect is recommended.")
    else:
        doc.add_paragraph("The player may benefit from dedicated training focusing on this specific ability. Consider fundamental drills and exercises.")
    
    # Add plot
    doc.add_heading('Player Comparison Plot', level=2)
    doc.add_picture(plot1, width=Inches(6))

    # Save the document to a BytesIO object
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# Function to plot comparison graph
def plot_player_comparison(new_player_score, score_df, metric_title):
    player_scores = score_df['Metric Value'].values
    player_names = score_df['Player'].values

    sorted_scores = np.sort(player_scores)
    sorted_indices = np.argsort(player_scores)
    sorted_names = player_names[sorted_indices]

    y_positions = np.arange(1, len(sorted_scores) + 1)

    fig, ax = plt.subplots(figsize=(12, 8))
    ax.plot(sorted_scores, y_positions, linestyle='-', color='blue', label='All Players')
    new_player_y = np.interp(new_player_score, sorted_scores, y_positions)

    ax.scatter(new_player_score, new_player_y, color='red', s=100, label='Your Score')
    ax.annotate("You're here", xy=(new_player_score, new_player_y), 
                xytext=(new_player_score + (max(sorted_scores) - min(sorted_scores))*0.05, new_player_y + 0.3),
                fontsize=12, color='darkred')

    ax.set_xlabel("Score Value", fontsize=12)
    ax.set_ylabel("Number of Players", fontsize=12)
    ax.set_title(f"Comparison of Your {metric_title} w.r.t Others", fontsize=16)
    ax.legend()
    ax.grid(True)

    plot_io = BytesIO()
    plt.savefig(plot_io, format='png')
    plot_io.seek(0)
    st.pyplot(fig)
    return plot_io

# Function to handle decision making
def handle_decision_making(folder_path):
    st.header("Decision Making Analysis")
    
    decision_metric = st.selectbox(
        "Select Decision Making Metric",
        ["Accuracy", "Response Time", "Decision Consistency"]
    )
    
    if st.button("Create Decision Making Metrics File"):
        avg_df, avg_file_path = calculate_metrics(folder_path, "Decision Making", output_folder=folder_path)
        if avg_file_path:
            st.success(f"Decision Making metrics file created at {avg_file_path}")
            st.dataframe(avg_df)
        else:
            st.error("Failed to create Decision Making metrics file")
    
    uploaded_file = st.file_uploader("Upload New Player's Decision Making Data", type="xlsx")
    if uploaded_file is not None:
        analyze_uploaded_file(uploaded_file, "Decision Making", folder_path, selected_metric=decision_metric)

# Function to handle stamina
def handle_stamina(folder_path):
    st.header("Stamina Analysis")

    stamina_option = st.radio(
        "Select Stamina Test:",
        ("Shuttle Run", "Other Stamina Test")
    )
    
    stamina_metric = st.selectbox(
        "Select Stamina Metric",
        ["Total Time", "Endurance Factor"]
    )
    
    if st.button("Create Stamina Metrics File"):
        avg_df, avg_file_path = calculate_metrics(folder_path, "Stamina", output_folder=folder_path)
        if avg_file_path:
            st.success(f"Stamina metrics file created at {avg_file_path}")
            st.dataframe(avg_df)
        else:
            st.error("Failed to create Stamina metrics file")

    uploaded_file = st.file_uploader(f"Upload New Player's {stamina_option} Data", type="xlsx")
    if uploaded_file is not None:
        analyze_uploaded_file(uploaded_file, "Stamina", folder_path, selected_metric=stamina_metric)

# Streamlit app
def app():
    st.title("Player Comparison Analysis")

    # Add AI info banner
    st.markdown("""
    <div style="background-color:rgba(255,159,67,0.2); padding:10px; border-radius:10px; margin-bottom:20px;">
    <h3 style="margin:0;color:#ff9f43">⚡ AI-Powered Analysis</h3>
    <p>This dashboard now includes personalized AI-powered recommendations for each player based on their performance metrics!</p>
    </div>
    """, unsafe_allow_html=True)

    metric_option = st.selectbox("Select Metric", ["Reaction Time", "Agility", "Decision Making", "Stamina"])

    if metric_option == "Reaction Time":
        handle_metric(reaction_time_folder_path, "Reaction Time")
    elif metric_option == "Agility":
        handle_metric(agility_folder_path, "Agility")
    elif metric_option == "Decision Making":
        handle_decision_making(decision_making_folder_path)
    elif metric_option == "Stamina":
        handle_stamina(stamina_folder_path)

# Run the app
if __name__ == "__main__":
    app()
                
